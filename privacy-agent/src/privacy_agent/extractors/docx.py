"""DOCX extractor via python-docx. Optional dependency."""
from __future__ import annotations

from pathlib import Path
from typing import Optional

import docx  # type: ignore[import-not-found]

from ..types import ExtractedContent
from .base import ExtractorError


class DocxExtractor:
    EXTENSIONS = (".docx",)

    def extract(self, path: Path, max_chars: Optional[int] = None) -> ExtractedContent:
        try:
            d = docx.Document(str(path))
        except Exception as e:
            raise ExtractorError(f"failed to open DOCX {path}: {e}") from e

        paragraphs = [p.text for p in d.paragraphs if p.text]
        text = "\n".join(paragraphs)

        if max_chars is not None and len(text) > max_chars:
            text = text[:max_chars]

        return ExtractedContent(
            text=text,
            title=path.stem,
            file_type="docx",
            line_count=len(paragraphs),
        )

    def extract_excerpt(
        self,
        path: Path,
        start_line: Optional[int] = None,
        end_line: Optional[int] = None,
        page: Optional[int] = None,
        max_chars: int = 500,
    ) -> str:
        d = docx.Document(str(path))
        paragraphs = [p.text for p in d.paragraphs if p.text]
        s = (start_line or 1) - 1
        e = end_line if end_line is not None else len(paragraphs)
        return "\n".join(paragraphs[s:e])[:max_chars]
